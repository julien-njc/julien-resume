#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from resume_model import load_resume


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "build" / "Julien_Pireaud_Resume_ATS.docx"


def set_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_line(document, text, bold_label=None, value=None, style=None, before=0, after=2):
    paragraph = document.add_paragraph(style=style)
    if text is not None:
        paragraph.add_run(text)
    else:
        if bold_label:
            run = paragraph.add_run(bold_label)
            run.bold = True
        paragraph.add_run(value or "")
    set_spacing(paragraph, before=before, after=after)
    return paragraph


def main():
    data = load_resume()
    OUTPUT.parent.mkdir(exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    title = document.add_paragraph()
    title_run = title.add_run(data["name"])
    title_run.bold = True
    title_run.font.size = Pt(18)
    set_spacing(title, after=8)

    add_line(document, None, "Address: ", data["address"])
    add_line(document, None, "Email: ", f"{data['email']} | Phone: {data['phone']}")
    add_line(document, None, "LinkedIn: ", data["linkedin"])
    add_line(document, None, "GitHub: ", data["github"])
    add_line(document, None, "Website: ", data["website"], after=6)

    def heading(text):
        p = document.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11.5)
        set_spacing(p, before=8, after=4)

    heading("Summary")
    add_line(document, data["summary"], after=4)

    heading("Technical Skills")
    for skill_name, skill_values in data["skills"].items():
        add_line(document, None, f"{skill_name}: ", ", ".join(skill_values))

    heading("Professional Experience")
    for role in data["experience"]:
        title = document.add_paragraph()
        title_run = title.add_run(f"{role['job_title']}, {role['employer']}")
        title_run.bold = True
        set_spacing(title, before=6, after=1)

        add_line(document, f"{role['location']} | {role['dates']}", after=2)
        for bullet in role["bullets"]:
            p = document.add_paragraph()
            p.add_run(f"- {bullet}")
            set_spacing(p, after=1)
        spacer = document.add_paragraph()
        set_spacing(spacer, after=3)

    heading("Education")
    for edu in data["education"]:
        add_line(document, None, "Degree: ", edu["degree"], after=1)
        add_line(document, None, "Field of Study: ", edu["field_of_study"], after=1)
        add_line(document, None, "School: ", edu["school"], after=1)
        add_line(document, None, "Graduation Date: ", edu["graduation_date"], after=4)

    heading("Languages")
    for language in data["languages"]:
        p = document.add_paragraph()
        p.add_run(f"- {language}")
        set_spacing(p, after=1)

    document.save(str(OUTPUT))


if __name__ == "__main__":
    main()
